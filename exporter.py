from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from datetime import datetime


def export_to_docx(answers_data: list, questionnaire_name: str = "Questionnaire") -> io.BytesIO:
    """
    Export answers to a DOCX file.
    
    Args:
        answers_data: List of answer tuples from database
                     Format: (id, questionnaire_id, question_number, question_text, 
                              generated_answer, edited_answer, citation, confidence, snippet)
        questionnaire_name: Name of the questionnaire for the
        
    Returns:
 title        BytesIO buffer containing the DOCX file
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading("SentraShield — Completed Security Questionnaire", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Questionnaire: {questionnaire_name}")
    doc.add_paragraph("")
    
    # Add summary section
    doc.add_heading("Summary", level=1)
    
    total_questions = len(answers_data)
    answered_with_citations = sum(1 for row in answers_data if row[4] != "Not found in references.")
    not_found_count = total_questions - answered_with_citations
    
    summary_text = f"Total Questions: {total_questions}\n"
    summary_text += f"Answered with Citations: {answered_with_citations}\n"
    summary_text += f"Not Found in References: {not_found_count}"
    doc.add_paragraph(summary_text)
    
    doc.add_paragraph("")
    # Add horizontal line using a paragraph with bottom border
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.border_bottom = True
    doc.add_paragraph("")
    
    # Add questions and answers
    doc.add_heading("Questions & Answers", level=1)
    
    for row in answers_data:
        # Unpack row data
        # row: (id, q_id, q_num, q_text, gen_answer, edited_answer, citation, confidence, snippet)
        q_num = row[2]
        q_text = row[3]
        gen_answer = row[4]
        edited_answer = row[5]
        citation = row[6]
        confidence = row[7]
        snippet = row[8] if len(row) > 8 else ""
        
        # Use edited answer if available, otherwise use generated
        answer = edited_answer if edited_answer else gen_answer
        
        # Question number and text
        p = doc.add_paragraph()
        run = p.add_run(f"Q{q_num}. ")
        run.bold = True
        run.font.size = Pt(12)
        
        run = p.add_run(q_text)
        run.bold = True
        run.font.size = Pt(11)
        
        # Answer
        answer_para = doc.add_paragraph()
        answer_run = answer_para.add_run("Answer: ")
        answer_run.bold = True
        answer_run.font.size = Pt(11)
        answer_para.add_run(answer)
        
        # Metadata: Citation and Confidence
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(f"Citation: {citation}  |  Confidence: {confidence:.2f}")
        meta_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        meta_run.font.size = Pt(9)
        
        # Evidence snippet (collapsible in Word, shown as indented text here)
        if snippet:
            snippet_para = doc.add_paragraph()
            snippet_run = snippet_para.add_run(f"Evidence: {snippet[:200]}...")
            snippet_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            snippet_run.font.size = Pt(8)
            snippet_para.paragraph_format.left_indent = Inches(0.5)
        
        # Add spacing between questions
        doc.add_paragraph("")
    
    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    
    return buf


def export_to_csv(answers_data: list) -> str:
    """
    Export answers to CSV format.
    
    Args:
        answers_data: List of answer tuples from database
        
    Returns:
        CSV formatted string
    """
    import csv
    import io
    
    output = io.StringIO()
    writer = csv.writer(output)
    
    # Write header
    writer.writerow([
        "Question Number", 
        "Question Text", 
        "Answer", 
        "Citation", 
        "Confidence",
        "Evidence Snippet"
    ])
    
    # Write data rows
    for row in answers_data:
        q_num = row[2]
        q_text = row[3]
        gen_answer = row[4]
        edited_answer = row[5]
        citation = row[6]
        confidence = row[7]
        snippet = row[8] if len(row) > 8 else ""
        
        # Use edited answer if available
        answer = edited_answer if edited_answer else gen_answer
        
        writer.writerow([
            q_num,
            q_text,
            answer,
            citation,
            confidence,
            snippet
        ])
    
    output.seek(0)
    return output.getvalue()
